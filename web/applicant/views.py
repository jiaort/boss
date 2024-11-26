# -*- coding:utf-8 -*-

from datetime import datetime

from django.db.models import Q, Count
from django.contrib import messages
from django.http import HttpResponse, HttpResponseRedirect
from django.urls import reverse
from docx import Document

from utils.paginater import paginate
from utils.errorcode import ERRORCODE
from utils.response import http_response
from utils.logger.syslogger import SysLogger
from web.decorators import render_admin, login_check, login_check_ajax

from app.applicant.models import Applicant
from app.applicant.constants import Source, Status
from app.job.constants import Education, WorkPlace, SubWorkPlace

from web.applicant.forms import ApplicantListForm


@login_check()
def list_view(request):

    form = ApplicantListForm(request.GET)
    form.is_valid()

    applicant_status_count = [0, 0, 0, 0, 0]
    applicant_queryset, total = paginate([], request.GET.get('page') or 1)
    data = {
        'active_classes': ['.applicant-menu'],
        'applicant_list': applicant_queryset,
        'total': total,
        'applicant_status_count': applicant_status_count,
        'status_option': Status.CHOICES,
        'form': form,
    }
    try:
        query = Q()
        name = form.cleaned_data.get('name')
        job_title = form.cleaned_data.get('job_title')
        status = form.cleaned_data.get('status')
        read = form.cleaned_data.get('read')
        if name:
            query &= Q(name__contains=name)
        if job_title is not None:
            query &= Q(job__job_title__contains=job_title)
        if status is not None:
            query &= Q(status=status)
        if read is not None:
            query &= Q(read=read)

        applicant_queryset = Applicant.objects.filter(query)

        status_order = form.cleaned_data.get('status_order')
        read_order = form.cleaned_data.get('read_order')
        if not status_order and not read_order:
            applicant_queryset = applicant_queryset.order_by('-id')

        if status_order == 1:
            applicant_queryset = applicant_queryset.order_by('status')
        elif status_order == 2:
            applicant_queryset = applicant_queryset.order_by('-status')

        if read_order == 1:
            applicant_queryset = applicant_queryset.order_by('read')
        elif read_order == 2:
            applicant_queryset = applicant_queryset.order_by('-read')

        applicant_queryset, total = paginate(applicant_queryset, request.GET.get('page') or 1)

        for item in applicant_queryset:
            item.source = Source.DICT.get(item.source)
            item.education = Education.DICT.get(item.education)
            item.work_place = WorkPlace.DICT.get(item.job.work_place) + '-' + SubWorkPlace.DICT.get(item.job.sub_work_place)
            item.status = Status.DICT.get(item.status)
            if item.file_applicant:
                item.file_applicant = item.file_applicant.split(',')

        data['applicant_list'] = applicant_queryset
        data['total'] = total

        applicant_annotate = Applicant.objects.values("status").annotate(count=Count("id"))

        total_status_count = 0
        for applicant in applicant_annotate:
            if applicant["status"] == Status.New:
                applicant_status_count[1] = applicant["count"]
            elif applicant["status"] == Status.Suitable:
                applicant_status_count[2] = applicant["count"]
            elif applicant["status"] == Status.Inappropriate:
                applicant_status_count[3] = applicant["count"]
            elif applicant["status"] == Status.Tobe:
                applicant_status_count[4] = applicant["count"]
            total_status_count += applicant["count"]

        applicant_status_count[0] = total_status_count

        data['applicant_status_count'] = applicant_status_count
    except Exception as ex:
        SysLogger.exception(ex, request)
        messages.error(request, '系统异常，请重试', 'danger')
        return render_admin(request, 'applicant/list.html', data)

    return render_admin(request, 'applicant/list.html', data)


@login_check()
def detail_view(request, a_id):
    applicant = Applicant.objects.filter(id=a_id).first()
    if not applicant:
        messages.error(request, u'应聘信息不存在')
        return HttpResponseRedirect(reverse('applicant_list', args=''))
    applicant.read = True
    applicant.save()
    data = {
        'active_classes': ['.applicant-menu'],
        "a_id": applicant.id,
        "name": applicant.name,
        "email": applicant.email,
        "mobile": applicant.mobile,
        "education": Education.DICT.get(applicant.education),
        "school": applicant.school,
        "profession": applicant.profession,
        "company": applicant.company,
        "position": applicant.position,
        "experience": applicant.experience,
        "applicant_list": applicant.file_applicant.split(",") if applicant.file_applicant else [],
        "status": applicant.status,
        "status_display": Status.DICT.get(applicant.status)
    }
    return render_admin(request, 'applicant/detail.html', data)


@login_check_ajax()
def change_status_view(request):
    applicant_id = request.POST.get('applicant_id')
    status = request.POST.get('status')

    Applicant.objects.filter(id=applicant_id).update(status=status)

    messages.error(request, '修改成功', 'danger')
    return http_response(status_code=ERRORCODE.SUCCESS)


@login_check_ajax()
def change_read_view(request):
    applicant_id = request.POST.get('applicant_id')
    applicant = Applicant.objects.filter(id=applicant_id).first()
    if applicant.read:
        applicant.read = False
    else:
        applicant.read = True
    applicant.save()

    messages.error(request, '修改成功', 'danger')
    return http_response(status_code=ERRORCODE.SUCCESS)


@login_check_ajax()
def unread_count_view(request):
    unread = Applicant.objects.filter(read=False).count()
    return http_response(status_code=ERRORCODE.SUCCESS, context={"unread": unread})


@login_check()
def export_view(request, applicant_id):
    """
    导出微简历
    :param request:
    :return:
    """
    applicant = Applicant.objects.filter(id=applicant_id).first()
    _now = datetime.now()
    response = HttpResponse(content_type='application/vnd.ms-word')
    response['Content-Disposition'] = 'attachment; filename=%04d%02d%02d%02d%02d%02d.docx' % \
                                      (_now.year, _now.month, _now.day, _now.hour, _now.minute, _now.second)
    document = Document()

    document.add_heading('个人简历', 0)

    document.add_heading('基本信息', level=5)

    name = document.add_paragraph()
    name.add_run('姓名：').bold = True
    name.add_run(applicant.name)

    mobile = document.add_paragraph()
    mobile.add_run('电话：').bold = True
    mobile.add_run(applicant.mobile)

    email = document.add_paragraph()
    email.add_run('邮箱：').bold = True
    email.add_run(applicant.email)

    document.add_heading('教育经历', level=5)

    education = document.add_paragraph()
    education.add_run('最高学历：').bold = True
    education.add_run(Education.DICT.get(applicant.education))

    school = document.add_paragraph()
    school.add_run('毕业院校：').bold = True
    school.add_run(applicant.school)

    profession = document.add_paragraph()
    profession.add_run('所学专业：').bold = True
    profession.add_run(applicant.profession)

    document.add_heading('工作经历', level=5)

    company = document.add_paragraph()
    company.add_run('当前公司：').bold = True
    company.add_run(applicant.company)

    position = document.add_paragraph()
    position.add_run('当前职位：').bold = True
    position.add_run(applicant.position)

    experience = document.add_paragraph()
    experience.add_run('总工作年限：').bold = True
    experience.add_run(applicant.experience)

    document.add_page_break()
    document.save(response)
    return response
